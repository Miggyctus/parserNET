import os
import json
import re

from datetime import datetime

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


REPORT_TEXT_PATH = "output/reports/llm_report.txt"
OUTPUT_DIR = "output/reports"
CHARTS_DIR = "output/charts"

os.makedirs(OUTPUT_DIR, exist_ok=True)

filename = f"{OUTPUT_DIR}/security_audit_report.docx"

doc = Document()


# =========================================================
# Helpers
# =========================================================

def clean_text(text):
    return str(text).replace("\n", " ").replace("\r", "").strip()


def insert_chart(chart_id):
    chart_path = os.path.join(CHARTS_DIR, f"{chart_id}.png")

    if os.path.exists(chart_path):
        try:
            doc.add_picture(chart_path, width=Inches(6.5))

            caption = doc.add_paragraph()
            caption.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            run = caption.add_run(
                chart_id.replace("_", " ").title()
            )
            run.bold = True

            os.remove(chart_path)
        except Exception as e:
            print(f"[WARN] Failed to insert chart {chart_id}: {e}")
    else:
        print(f"[WARN] Chart not found: {chart_path}")


# =========================================================
# Markdown Parsing
# =========================================================

def add_heading_from_markdown(line):
    level = line.count("#")

    text = line.replace("#", "").strip()

    level = max(1, min(level, 3))

    doc.add_heading(text, level=level)


def add_paragraph_with_bold(text):
    paragraph = doc.add_paragraph()

    parts = re.split(r"(\*\*.*?\*\*)", text)

    for part in parts:

        if part.startswith("**") and part.endswith("**"):
            clean = part[2:-2]

            run = paragraph.add_run(clean)
            run.bold = True

        else:
            paragraph.add_run(part)


def add_unordered_list(lines):
    for line in lines:

        text = re.sub(r"^[\-\*]\s+", "", line.strip())

        doc.add_paragraph(
            clean_text(text),
            style="List Bullet"
        )


def add_ordered_list(lines):
    for line in lines:

        text = re.sub(r"^\d+\.\s+", "", line.strip())

        doc.add_paragraph(
            clean_text(text),
            style="List Number"
        )


# =========================================================
# Robust Markdown Table Parser
# =========================================================

def add_table_from_markdown(table_lines):

    rows = []

    for line in table_lines:

        line = line.strip()

        if not line:
            continue

        # Skip markdown separator rows
        if re.match(r'^\|?[\-\:\s\|]+\|?$', line):
            continue

        if "|" not in line:
            continue

        parts = [
            clean_text(cell)
            for cell in line.strip("|").split("|")
        ]

        # Remove trailing empty cells
        while parts and parts[-1] == "":
            parts.pop()

        rows.append(parts)

    if not rows:
        return

    # =====================================================
    # Normalize rows
    # =====================================================

    max_cols = max(len(r) for r in rows)

    normalized_rows = []

    for row in rows:

        # Fill missing cols
        if len(row) < max_cols:
            row.extend([""] * (max_cols - len(row)))

        # Trim overflow
        normalized_rows.append(row[:max_cols])

    # =====================================================
    # Create Word table
    # =====================================================

    table = doc.add_table(
        rows=len(normalized_rows),
        cols=max_cols
    )

    table.style = "Table Grid"

    for row_idx, row in enumerate(normalized_rows):

        for col_idx, cell in enumerate(row):

            try:
                table.rows[row_idx].cells[col_idx].text = clean_text(cell)

            except Exception as e:
                print(
                    f"[WARN] Failed table cell "
                    f"({row_idx},{col_idx}): {e}"
                )


# =========================================================
# Main Report Parsing
# =========================================================

if os.path.exists(REPORT_TEXT_PATH):

    with open(REPORT_TEXT_PATH, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0

    while i < len(lines):

        line = lines[i].rstrip()

        # =================================================
        # Empty line
        # =================================================

        if not line.strip():
            i += 1
            continue

        # =================================================
        # Headings
        # =================================================

        if line.strip().startswith("#"):

            add_heading_from_markdown(line)

            i += 1
            continue

        # =================================================
        # Tables
        # =================================================

        if "|" in line:

            table_block = []

            while i < len(lines) and "|" in lines[i]:

                table_block.append(lines[i])

                i += 1

            add_table_from_markdown(table_block)

            continue

        # =================================================
        # Horizontal Rules
        # =================================================

        if line.strip().startswith("---"):

            i += 1
            continue

        # =================================================
        # Chart Placeholders
        # =================================================

        placeholder_match = re.match(
            r"\{\{CHART:\s*([a-z0-9_]+).*?\}\}",
            line.strip()
        )

        if placeholder_match:

            chart_id = placeholder_match.group(1)

            insert_chart(chart_id)

            i += 1
            continue

        # =================================================
        # Ordered List
        # =================================================

        if re.match(r"^\d+\.\s+", line.strip()):

            list_block = []

            while (
                i < len(lines)
                and re.match(r"^\d+\.\s+", lines[i].strip())
            ):
                list_block.append(lines[i])
                i += 1

            add_ordered_list(list_block)

            continue

        # =================================================
        # Unordered List
        # =================================================

        if re.match(r"^[\-\*]\s+", line.strip()):

            list_block = []

            while (
                i < len(lines)
                and re.match(r"^[\-\*]\s+", lines[i].strip())
            ):
                list_block.append(lines[i])
                i += 1

            add_unordered_list(list_block)

            continue

        # =================================================
        # Normal Paragraph
        # =================================================

        add_paragraph_with_bold(clean_text(line))

        i += 1

else:

    doc.add_paragraph("Report text not found.")


# =========================================================
# Save
# =========================================================

doc.save(filename)

print(json.dumps({
    "generated_report": filename
}))